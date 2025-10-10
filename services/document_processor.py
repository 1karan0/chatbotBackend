from typing import List
import io
import re
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from config.settings import settings

try:
    import nltk
    nltk.download('punkt', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
except Exception as e:
    print(f"Warning: NLTK download failed: {e}")

class DocumentProcessor:
    """Handles document processing, chunking, and embedding generation."""

    def __init__(self):
        self.embeddings = OpenAIEmbeddings(model=settings.embedding_model)
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        self.batch_size = 100  # number of chunks per embedding request

    def sanitize_text(self, text: str) -> str:
        """Remove null bytes to prevent Postgres errors."""
        return re.sub(r"[\x00]", "", text)

    def process_text(self, text: str, source: str, tenant_id: str) -> List[Document]:
        """Split text into chunks and attach metadata."""
        text = self.sanitize_text(text)
        documents = [Document(
            page_content=text,
            metadata={"source": source, "tenant_id": tenant_id}
        )]
        chunks = self.text_splitter.split_documents(documents)
        return chunks

    async def generate_embeddings(self, chunks: List[Document]) -> List[List[float]]:
        """Generate embeddings for document chunks in safe batches."""
        all_texts = [c.page_content for c in chunks]
        embeddings_list = []

        for i in range(0, len(all_texts), self.batch_size):
            batch = all_texts[i:i + self.batch_size]
            try:
                batch_embeddings = await self.embeddings.aembed_documents(batch)
                embeddings_list.extend(batch_embeddings)
            except Exception as e:
                print(f"Error generating embeddings for batch {i}-{i+len(batch)}: {e}")

        return embeddings_list

    def extract_file_content(self, file_content: bytes, file_name: str) -> str:
        """Extract text content from file bytes and sanitize."""
        try:
            if file_name.endswith(('.txt', '.md', '.csv')):
                text = file_content.decode('utf-8', errors='ignore')
            elif file_name.endswith('.pdf'):
                try:
                    from pypdf import PdfReader
                    pdf_file = io.BytesIO(file_content)
                    pdf_reader = PdfReader(pdf_file)
                    text = "\n".join([page.extract_text() or "" for page in pdf_reader.pages])
                except Exception as e:
                    print(f"Error extracting PDF: {e}")
                    text = file_content.decode('utf-8', errors='ignore')
            elif file_name.endswith('.docx'):
                try:
                    from docx import Document as DocxDocument
                    doc_file = io.BytesIO(file_content)
                    doc = DocxDocument(doc_file)
                    text = "\n".join([p.text for p in doc.paragraphs])
                except Exception as e:
                    print(f"Error extracting DOCX: {e}")
                    text = file_content.decode('utf-8', errors='ignore')
            else:
                text = file_content.decode('utf-8', errors='ignore')

            return self.sanitize_text(text)

        except Exception as e:
            print(f"Error extracting file content: {e}")
            return ""

document_processor = DocumentProcessor()
